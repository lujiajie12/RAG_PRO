from __future__ import annotations

from io import BytesIO

from docx import Document

from .base_loader import BaseLoader
from ..types import ParsedDocument


class DocxLoader(BaseLoader):
    def parse_bytes(self, data: bytes, file_name: str) -> ParsedDocument:
        document = Document(BytesIO(data))
        segments = []

        for paragraph in document.paragraphs:
            text = " ".join(paragraph.text.split())
            if text:
                kind = "heading" if paragraph.style and paragraph.style.name.lower().startswith("heading") else "paragraph"
                level = None
                if kind == "heading":
                    suffix = paragraph.style.name.split()[-1]
                    level = int(suffix) if suffix.isdigit() else None
                segment = self.make_segment(kind, text, style=paragraph.style.name if paragraph.style else "")
                if level is not None:
                    segment.metadata["heading_level"] = level
                segments.append(segment)

        table_rows = 0
        for table in document.tables:
            for row in table.rows:
                cell_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cell_text:
                    table_rows += 1
                    segments.append(self.make_segment("table_row", " | ".join(cell_text), cell_count=len(cell_text)))

        return ParsedDocument(
            file_name=file_name,
            file_type="docx",
            parser_name="docx",
            segments=segments,
            metadata={"paragraph_count": len(document.paragraphs), "table_row_count": table_rows},
        )
